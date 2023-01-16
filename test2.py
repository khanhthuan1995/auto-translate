from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import requests
import re
import click

api_url = "http://nlpapi.h5online.xyz:6060/translate"



def translate(data):
    params = {
        "text": data,
        "source_lang": "en",
        "target_lang": "vi"
    }
    response = requests.post(api_url, json=params)
    dataJson = response.json()
    return dataJson['result']['script']
# dem = 0
# p = document.paragraphs[29]
# for r in p.runs:
#     print(r.font.bold)

def checkAlphaNum(text):
    return len(text)>1 or re.match('^[a-zA-Z0-9_ÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚĂĐĨŨƠàáâãèéêìíòóôõùúăđĩũơƯĂẠẢẤẦẨẪẬẮẰẲẴẶẸẺẼỀỀỂưăạảấầẩẫậắằẳẵặẹẻẽềềểỄỆỈỊỌỎỐỒỔỖỘỚỜỞỠỢỤỦỨỪễệỉịọỏốồổỗộớờởỡợụủứừỬỮỰỲỴÝỶỸửữựỳỵỷỹ,.]+$' , text)

def checkDifferent(run1, run2, k):
    return (not (len(run1.text)==1 and run2.text != ' ' and k ==1)) and (run1.font.color.rgb != run2.font.color.rgb or run1.font.size != run2.font.size or run1.font.name != run2.font.name or (run1.font.bold == True and run2.font.bold == None) or run1.underline != run2.underline)

def getText(runs, k, n):
    if k == n-1:
        return {
            "k": k+1,
            "text": runs[k].text.strip()
        }
    text = runs[k].text
    k = k + 1
    if not checkAlphaNum(text):
        return {
            "k": k,
            "text": text.strip()
        }
    for i in range(k, n):
        if checkDifferent(runs[i-1], runs[i], i):
            return {
                "k": i,
                "text": text.strip(),
            }
        elif not checkAlphaNum(runs[i].text):
            return {
                "k": i,
                "text": text.strip()
            }
        else:
            text = text + runs[i].text
    return {
        "k": n,
        "text": text.strip()
    }

# p = document.paragraphs[29]
# k = 0
# n = len(p.runs)
# print(p.runs[6].text==" ")
# for r  in p.runs:
#     print(r.text)
# if n != 0:
#     data = getText(p.runs, k, n)
#     if data['text']:
#         print(data['text'], translate(data['text']))
#         data['text'] = translate(data['text'])

#     newK = data['k']
#     p.runs[k].text = data['text']
#     k = k+1
#     for i in range(k, newK):
#         p.runs[i].text = ''
#     k = newK

#     while k < n:
#         data = getText(p.runs, k, n)
#         if data['text']:
#             print(data['text'], translate(data['text']))
#             data['text'] = translate(data['text'])
#         newK = data['k']
#         p.runs[k].text = data['text']
#         k = k+1
#         for i in range(k, newK):
#             p.runs[i].text = ''
#         k = newK 



def transParagraphs(p):
    # print(p.text)
    k = 0
    n = len(p.runs)
    if n != 0:
        data = getText(p.runs, k, n)
        if checkAlphaNum(data['text']):
            # print(not checkDifferent(p.runs[0], p.runs[1], 1),p.runs[0].text, p.runs[1].text)
            # print(not checkAlphaNum(p.runs[1].text), p.runs[1].text, len(p.runs[1].text))
            textTranslate = translate(data['text'])
            # print(data['text'], textTranslate)
            data['text'] = textTranslate
            if n >= 2 and len(data['text']) > 1 and not checkDifferent(p.runs[0], p.runs[1], 1) and data['k'] > 1:
                p.runs[k].text = data['text'][0]
                p.runs[k+1].text = data['text'][1:len(data['text'])].replace('\t', ' ').replace('•', '')
                k = k+2
                newK = data['k']
                for i in range(k, newK):
                    p.runs[i].text = ''
                k = newK
            else:
                newK = data['k']
                p.runs[k].text = data['text'].replace('\t', ' ').replace('•', '')
                k = k+1
                for i in range(k, newK):
                    p.runs[i].text = ''
                k = newK
        else:
            k = k+1

        while k < n:
            data = getText(p.runs, k, n)
            if checkAlphaNum(data['text']):
                textTranslate = translate(data['text'])
                # print(data['text'], textTranslate)
                data['text'] = textTranslate
                newK = data['k']
                p.runs[k].text = data['text'].replace('\t', ' ').replace('•', '')
                k = k+1
                for i in range(k, newK):
                    p.runs[i].text = ''
                k = newK 
            else:
                k = k+1


@click.command(context_settings=dict(ignore_unknown_options=True))
@click.option('--input', prompt='Nhập tên tệp cần dịch', help='Tên tệp có chứa cả .docx nhé')
@click.option('--output', prompt='Nhập tên sẽ được lưu',
              help='Tên mới cần lưu',)
def docprocessing(input, output):
    document = Document(input)
    click.echo(click.style(f"Translating...", blink=True, fg="red", bg='blue'))
    with click.progressbar(document.paragraphs) as bar:
        for p in bar:
            transParagraphs(p)
        
        for t in document.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        transParagraphs(p)
    filename = output+".docx"
    document.save(filename)   
    click.echo(click.style(f"Completed, Now you can access translated document", bold=True, fg="green"))
    click.echo('Path: %s' % click.format_filename(filename))
# rels = document.part.rels
# for rel in rels:
#     if rels[rel].reltype == RT.HYPERLINK:
#         print(rels[rel]._target)
        # rels[rel]._target = new_url


if __name__ == "__main__":
    docprocessing()