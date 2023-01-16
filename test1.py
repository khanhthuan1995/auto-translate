from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import requests
api_url = "http://nlpapi.h5online.xyz:6060/translate"
document = Document('Test1.docx')


def translate(data):
    params = {
        "text": data,
        "source_lang": "en",
        "target_lang": "vi"
    }
    response = requests.post(api_url, json=params)
    dataJson = response.json()
    return dataJson['result']['script']
# dem = 0
# p = document.paragraphs[29]
# for r in p.runs:
#     print(r.font.name)


def getText(runs, k, n):
    if k == n-1:
        return {
            "k": k+1,
            "text": runs[k].text.strip()
        }
    text = runs[k].text
    k = k + 1
    if text == ' ':
        return {
            "k": k,
            "text": ' '
        }
    for i in range(k, n):
        if runs[i].font.color.rgb != runs[i-1].font.color.rgb or runs[i].font.size != runs[i-1].font.size or runs[i].font.name != runs[i-1].font.name:
            return {
                "k": i,
                "text": text.strip()
            }
        elif runs[i].text == ' ':
            return {
                "k": i,
                "text": text.strip()
            }
        else:
            text = text + runs[i].text
    return {
        "k": n,
        "text": text.strip()
    }

# p = document.paragraphs[29]
# k = 0
# n = len(p.runs)
# print(p.runs[6].text==" ")
# for r  in p.runs:
#     print(r.text)
# if n != 0:
#     data = getText(p.runs, k, n)
#     if data['text']:
#         print(data['text'], translate(data['text']))
#         data['text'] = translate(data['text'])

#     newK = data['k']
#     p.runs[k].text = data['text']
#     k = k+1
#     for i in range(k, newK):
#         p.runs[i].text = ''
#     k = newK

#     while k < n:
#         data = getText(p.runs, k, n)
#         if data['text']:
#             print(data['text'], translate(data['text']))
#             data['text'] = translate(data['text'])
#         newK = data['k']
#         p.runs[k].text = data['text']
#         k = k+1
#         for i in range(k, newK):
#             p.runs[i].text = ''
#         k = newK 



def transParagraphs(p):
    k = 0
    n = len(p.runs)
    if n != 0:
        data = getText(p.runs, k, n)
        if data['text'].strip():
            print(data['text'], translate(data['text']))
            data['text'] = translate(data['text'].strip())
            newK = data['k']
            p.runs[k].text = data['text']
            k = k+1
            for i in range(k, newK):
                p.runs[i].text = ''
            k = newK
        else:
            k = k+1

        while k < n:
            data = getText(p.runs, k, n)
            if data['text'].strip():
                print(data['text'], translate(data['text']))
                data['text'] = translate(data['text'].strip())
                newK = data['k']
                p.runs[k].text = data['text']
                k = k+1
                for i in range(k, newK):
                    p.runs[i].text = ''
                k = newK 
            else:
                k = k+1


for p in document.paragraphs:
    transParagraphs(p)

for t in document.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                transParagraphs(p)

# rels = document.part.rels
# for rel in rels:
#     if rels[rel].reltype == RT.HYPERLINK:
#         print(rels[rel]._target)
        # rels[rel]._target = new_url

document.save('output1.docx')

print('done')
